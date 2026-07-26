"""Procesador de documentos Word para extracción de texto."""
import os
from typing import List, Dict
from docx import Document
from .config import Config

class DocumentProcessor:
    """Procesa archivos Word y los prepara para indexación."""
    
    def __init__(self, config: Config):
        self.config = config
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extrae texto de un archivo .docx."""
        doc = Document(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n".join(text_parts)
    
    def chunk_text(self, text: str) -> List[str]:
        """Divide el texto en chunks de tamaño manejable."""
        paragraphs = text.split('\n')
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.config.CHUNK_SIZE:
                current_chunk += "\n" + para if current_chunk else para
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def process_documents(self, data_dir: str = "data") -> List[Dict]:
        """Procesa todos los documentos Word en el directorio data/."""
        documents = []
        
        for filename in os.listdir(data_dir):
            if filename.endswith(".docx"):
                file_path = os.path.join(data_dir, filename)
                category_key = filename.replace(".docx", "")
                category_name = self.config.DOCUMENT_MAPPING.get(
                    category_key, category_key.title()
                )
                
                text = self.extract_text_from_docx(file_path)
                chunks = self.chunk_text(text)
                
                for i, chunk in enumerate(chunks):
                    documents.append({
                        "page_content": chunk,
                        "metadata": {
                            "source": filename,
                            "category": category_name,
                            "chunk_id": i,
                            "file_path": file_path
                        }
                    })
        
        return documents

      
          
